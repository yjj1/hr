# -*- coding:'utf8' -*-
#encoding=utf-8
from xml.dom import minidom

import docx
import zipfile
import types
import uuid
from orm import *
from connect_db import ConnectDb
from win32com import client as wc
import os
#global var
global connectDb

class ParseDoc:
    emailId = None
    def __init__(self):
        print 'init'

    def setEmailId(self, eId):
        self.emailId = eId

    def parse(self,file_path):
        path_array = os.path.splitext(file_path)
        file_name = path_array[0]
        suffix = path_array[1]
        sDoc = '.doc'
        if suffix == sDoc:
            file_path = self.doc2docx(file_path)
        fullText = []
        doc = docx.Document(file_path)
        tables = doc.tables

        # for i_r, row in enumerate(tables[0].rows):
        #     tmp_i = -1
        #     for cell in row.cells:
        #         tmp_i = tmp_i + 1
        #         cell_data = []
        #         # print cell.length
        #         for p in cell.paragraphs:
        #             cell_data.append(p.text)
        #             print p.text
        #             fullText.append(p)
        t = tables[0]
        for i in range(0, 330):
            content = t.cell(0, i).text
            print content, '(',0,',',i, ')'

        sUuid = str(uuid.uuid4())
        data = ResumeBaseInfo(sUuid)

        name = t.cell(0,1).text
        sex = t.cell(0,5).text
        birthday = t.cell(0,11).text
        eduBack = t.cell(0,17).text
        height = t.cell(0, 21).text
        health = t.cell(0, 27).text
        politicsStatus = t.cell(0, 33).text
        nativePlace = t.cell(0,37).text
        marriage = t.cell(0,43).text
        address = t.cell(0,49).text
        idCard = t.cell(0,65).text
        phone = t.cell(0,84).text
        homePhone = t.cell(0,93).text
        czSchool = t.cell(0,114).text
        czPro = t.cell(0,121).text
        czDate = t.cell(0,124).text

        gzSchool = t.cell(0,130).text
        gzPro = t.cell(0,137).text
        gzDate = t.cell(0,140).text

        dzSchool = t.cell(0,146).text
        dzPro = t.cell(0,153).text
        dzDate = t.cell(0,156).text

        bkSchool = t.cell(0,162).text
        bkPro = t.cell(0,169).text
        bkDate = t.cell(0,172).text

        yjsSchool = t.cell(0, 178).text
        yjsPro = t.cell(0, 185).text
        yjsDate = t.cell(0, 188).text

        bigBang = t.cell(0,193).text

        job = t.cell(0,305).text
        isAddressControl = t.cell(0,327).text

        #家庭
        f1 = Family(str(uuid.uuid4()))
        f1.resumeId = data.resumeId
        f1.name=t.cell(0,225).text
        f1.relationShip=t.cell(0,227).text
        f1.politicsStatus=t.cell(0,230).text
        f1.company=t.cell(0,233).text
        f1.job=t.cell(0,239).text

        f2 = Family(str(uuid.uuid4()))
        f2.resumeId = data.resumeId
        f2.name = t.cell(0, 241).text
        f2.relationShip = t.cell(0, 243).text
        f2.politicsStatus = t.cell(0, 246).text
        f2.company = t.cell(0, 249).text
        f2.job = t.cell(0, 255).text

        f3 = Family(str(uuid.uuid4()))
        f3.resumeId = data.resumeId
        f3.name = t.cell(0, 257).text
        f3.relationShip = t.cell(0, 259).text
        f3.politicsStatus = t.cell(0, 262).text
        f3.company = t.cell(0, 265).text
        f3.job = t.cell(0, 271).text

        f4 = Family(str(uuid.uuid4()))
        f4.resumeId = data.resumeId
        f4.name = t.cell(0, 273).text
        f4.relationShip = t.cell(0, 275).text
        f4.politicsStatus = t.cell(0, 278).text
        f4.company = t.cell(0, 281).text
        f4.job = t.cell(0, 287).text

        data.name = name
        data.sex = sex
        data.birthday = birthday
        data.eduBack = eduBack
        data.phone = phone
        data.homePhone = homePhone
        data.resumePath = file_path
        data.politicsStatus = politicsStatus
        data.nativePlace = nativePlace
        data.address = address
        data.bigBang = bigBang
        data.height = height
        data.healthy = health
        data.marriage = marriage
        data.idCard = idCard
        data.job = job
        data.isAddressControl = isAddressControl

        eduDataCz = Edu(str(uuid.uuid4()))
        eduDataCz.resumeId = data.resumeId
        eduDataCz.eduDate = czDate
        eduDataCz.eduPro = czPro
        eduDataCz.eduType = '1'
        eduDataCz.schoolName = czSchool

        eduDataGz = Edu(str(uuid.uuid4()))
        eduDataGz.resumeId = data.resumeId
        eduDataGz.eduDate = gzDate
        eduDataGz.eduPro = gzPro
        eduDataGz.eduType = '2'
        eduDataGz.schoolName = gzSchool

        eduDataDz = Edu(str(uuid.uuid4()))
        eduDataDz.resumeId = data.resumeId
        eduDataDz.eduDate = dzDate
        eduDataDz.eduPro = dzPro
        eduDataDz.eduType = '3'
        eduDataDz.schoolName = dzSchool

        eduDataBk = Edu(str(uuid.uuid4()))
        eduDataBk.resumeId = data.resumeId
        eduDataBk.eduDate = bkDate
        eduDataBk.eduPro = bkPro
        eduDataBk.eduType = '4'
        eduDataBk.schoolName = bkSchool

        eduDataYjs = Edu(str(uuid.uuid4()))
        eduDataYjs.resumeId = data.resumeId
        eduDataYjs.eduDate = yjsDate
        eduDataYjs.eduPro = yjsPro
        eduDataYjs.eduType = '5'
        eduDataYjs.schoolName = yjsSchool

        connectDb = ConnectDb()
        session = connectDb.db_session
        session.add(data)
        session.add(f1)
        session.add(f2)
        session.add(f3)
        session.add(f4)
        session.add(eduDataCz)
        session.add(eduDataGz)
        session.add(eduDataDz)
        session.add(eduDataBk)
        session.add(eduDataYjs)

        session.commit()

        # print t.cell(0,8).text
        # print t.cell(0,9).text
        # print t.cell(0,10).text
        # print t.cell(0,11).text

        #如果从邮件中来，则更新邮件状态
        if self.emailId is not None:
            session.query(OriginEmail).filter(OriginEmail.emailId == self.emailId).update({
                OriginEmail.status:'1'
            })
            session.commit()

        return fullText

    def zipParse(self, file_path):
        document = zipfile.ZipFile(file_path)
        xml_content = document.read('word/document.xml')
        reparsed = minidom.parseString(xml_content)
        print reparsed.toprettyxml(indent="   ", encoding="utf-8")

    def doc2docx(self, file_path):
        new_file_path = file_path + 'x'

        word = wc.Dispatch('Word.Application')
        doc = word.Documents.Open(file_path)
        doc.SaveAs(new_file_path, FileFormat=16)
        doc.Close()

        word.Quit()

        return new_file_path

if __name__ == '__main__':
    text = ParseDoc().parse('D:\\test2.doc')
    # ParseDoc().zipParse('/users/jianjingye/test2.doc')
    # for t in text:
    #     print t.text